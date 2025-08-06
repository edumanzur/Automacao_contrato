# Native packages
import os
import re
import base64
import shutil
import tempfile
from datetime import datetime

# Imported packages
from fastapi import HTTPException, APIRouter
from fastapi.responses import Response
from docx import Document

# local packages
from app.services import document_processor
from app.core import logger


router = APIRouter(prefix="/documents")


@router.get("/")
def root():
    return {
        "message": "N8N + WhatsApp Cloud API",
        "version": "2.0.2",
        "status": "online",
        "environment": "production",
        "timestamp": "datetime.now().isoformat()",
        "endpoints": {
            "generate_document": "POST /generate-document (returns binary)",
            "generate_document_base64": "POST /generate-document-base64 (returns JSON with base64)",
            "generate_document_whatsapp": "POST /generate-document-whatsapp (optimized for Z-API)",
            "webhook": "POST /webhook/process",
            "health": "GET /health",
            "test_substitution": "POST /test-substitution (for debug)"
        }
    }


# @app.post("/test-substituicao")
# async def test_substituicao():
#     """Endpoint para testar substituições de placeholder"""
#     dados_teste = {
#         "NOME": "João Silva",
#         "VALOR": "R$ 1.500,00",
#         "EMAIL": "joao@email.com",
#         "CPF": "123.456.789-00",
#         "TELEFONE": "(11) 99999-9999"
#     }
    
#     # Teste simples de string
#     texto_teste = "Nome: {{NOME}}, Valor: {{VALOR}}, Email: {{EMAIL}}"
#     resultado_string = texto_teste
    
#     for chave, valor in dados_teste.items():
#         placeholder = f'{{{{{chave}}}}}'
#         if placeholder in resultado_string:
#             resultado_string = resultado_string.replace(placeholder, str(valor))
    
#     return {
#         "success": True,
#         "teste_string": {
#             "original": texto_teste,
#             "resultado": resultado_string
#         },
#         "dados_teste": dados_teste,
#         "timestamp": datetime.now().isoformat()
#     }


# @app.post("/gerar-documento")
# async def gerar_documento(request: MensagemRequest):
#     """Endpoint para processar mensagem E gerar documento DOCX (retorna binário)"""
#     logger.info("=== GERAÇÃO DE DOCUMENTO N8N CLOUD (BINÁRIO) ===")
#     logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
#     temp_dir = tempfile.mkdtemp()
#     output_path = None
    
#     try:
#         # Extrair dados
#         dados_extraidos = services.extrair_dados_da_mensagem(request.mensagem)
#         logger.info("Dados extraídos para documento")
        
#         # Definir nome do arquivo
#         nome_cliente = dados_extraidos.get("NOME", "cliente").replace(" ", "_")
#         nome_cliente = re.sub(r'[^\w\-_.]', '', nome_cliente)
#         timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
#         output_filename = f"documento_{nome_cliente}_{timestamp}.docx"
#         output_path = os.path.join(temp_dir, output_filename)
        
#         logger.info(f"Gerando documento: {output_filename}")
        
#         try:
#             # Procurar template
#             possible_templates = [
#                 DocxNames.TEMPLATE,
#                 DocxNames.MODELO,
#                 "templates/" + DocxNames.TEMPLATE,
#                 "templates/" + DocxNames.MODELO
#             ]
            
#             template_encontrado = None
#             for template_path in possible_templates:
#                 if os.path.exists(template_path):
#                     template_encontrado = template_path
#                     break
            
#             if template_encontrado:
#                 logger.info(f"Template encontrado: {template_encontrado}")
#                 services.preencher_modelo(template_encontrado, output_path, dados_extraidos)
#                 logger.info("Template preenchido com sucesso")
#             else:
#                 logger.info("Template não encontrado, criando documento padrão")
#                 services.criar_documento_fallback(dados_extraidos, output_path)
                
#         except Exception as e:
#             logger.error(f"Erro no preenchimento: {e}")
#             logger.info("Criando documento fallback...")
#             services.criar_documento_fallback(dados_extraidos, output_path)
        
#         # Verificar se arquivo foi criado
#         if not os.path.exists(output_path):
#             raise Exception("Documento não foi gerado")
        
#         file_size = os.path.getsize(output_path)
#         logger.info(f"Documento criado: {file_size} bytes")
        
#         # Ler arquivo
#         with open(output_path, "rb") as f:
#             docx_content = f.read()
        
#         return Response(
#             content=docx_content,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={
#                 "Content-Disposition": f"attachment; filename={output_filename}",
#                 "Content-Length": str(len(docx_content)),
#                 "Cache-Control": "no-cache",
#                 "X-Filename": output_filename,
#                 "X-File-Size": str(file_size)
#             }
#         )
        
#     except Exception as e:
#         logger.error(f"ERRO CRÍTICO: {e}")
#         raise HTTPException(status_code=500, detail=f"Erro na geração do documento: {str(e)}")
    
#     finally:
#         # Limpar arquivos temporários
#         try:
#             if temp_dir and os.path.exists(temp_dir):
#                 shutil.rmtree(temp_dir)
#                 logger.info("Arquivos temporários removidos")
#         except Exception as e:
#             logger.warning(f"Erro na limpeza: {e}")


# @app.post("/gerar-documento-base64", response_model=DocumentoResponse)
# async def gerar_documento_base64(request: MensagemRequest):
#     """Endpoint que retorna o documento em base64 (ideal para integração com APIs)"""
#     logger.info("=== GERAÇÃO DE DOCUMENTO N8N CLOUD (BASE64) ===")
#     logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
#     temp_dir = tempfile.mkdtemp()
#     output_path = None
    
#     try:
#         # Extrair dados
#         dados_extraidos = services.extrair_dados_da_mensagem(request.mensagem)
#         logger.info("Dados extraídos para documento")
        
#         # Definir nome do arquivo
#         nome_cliente = dados_extraidos.get("NOME", "cliente").replace(" ", "_")
#         nome_cliente = re.sub(r'[^\w\-_.]', '', nome_cliente)
#         timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
#         output_filename = f"documento_{nome_cliente}_{timestamp}.docx"
#         output_path = os.path.join(temp_dir, output_filename)
        
#         logger.info(f"Gerando documento: {output_filename}")
        
#         try:
#             # Procurar template
#             possible_templates = [
#                 "template.docx",
#                 "modelo.docx",
#                 "templates/template.docx",
#                 "templates/modelo.docx"
#             ]
            
#             template_encontrado = None
#             for template_path in possible_templates:
#                 if os.path.exists(template_path):
#                     template_encontrado = template_path
#                     break
            
#             if template_encontrado:
#                 logger.info(f"Template encontrado: {template_encontrado}")
#                 services.preencher_modelo(template_encontrado, output_path, dados_extraidos)
#                 logger.info("Template preenchido com sucesso")
#             else:
#                 logger.info("Template não encontrado, criando documento padrão")
#                 services.criar_documento_fallback(dados_extraidos, output_path)
                
#         except Exception as e:
#             logger.error(f"Erro no preenchimento: {e}")
#             logger.info("Criando documento fallback...")
#             services.criar_documento_fallback(dados_extraidos, output_path)
        
#         # Verificar se arquivo foi criado
#         if not os.path.exists(output_path):
#             raise Exception("Documento não foi gerado")
        
#         file_size = os.path.getsize(output_path)
#         logger.info(f"Documento criado: {file_size} bytes")
        
#         # Converter para base64
#         with open(output_path, "rb") as f:
#             docx_content = f.read()
#             base64_content = base64.b64encode(docx_content).decode('utf-8')
        
#         mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
#         return DocumentoResponse(
#             success=True,
#             message="Documento gerado com sucesso",
#             filename=output_filename,
#             file_size=file_size,
#             mime_type=mime_type,
#             base64_content=base64_content,
#             dados_extraidos=dados_extraidos,
#             timestamp=datetime.now().isoformat()
#         )
        
#     except Exception as e:
#         logger.error(f"ERRO CRÍTICO: {e}")
#         raise HTTPException(status_code=500, detail=f"Erro na geração do documento: {str(e)}")
    
#     finally:
#         # Limpar arquivos temporários
#         try:
#             if temp_dir and os.path.exists(temp_dir):
#                 shutil.rmtree(temp_dir)
#                 logger.info("Arquivos temporários removidos")
#         except Exception as e:
#             logger.warning(f"Erro na limpeza: {e}")


# @app.post("/gerar-documento-whatsapp")
# async def gerar_documento_whatsapp(request: MensagemRequest):
#     """Endpoint otimizado para envio via WhatsApp usando Z-API"""
#     logger.info("=== GERAÇÃO DE DOCUMENTO PARA WHATSAPP ===")
#     logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
#     temp_dir = tempfile.mkdtemp()
#     output_path = None
    
#     try:
#         # Extrair dados
#         dados_extraidos = services.extrair_dados_da_mensagem(request.mensagem)
#         logger.info("Dados extraídos para documento WhatsApp")
        
#         # Definir nome do arquivo (mais curto para WhatsApp)
#         nome_cliente = dados_extraidos.get("NOME", "cliente").replace(" ", "_")
#         nome_cliente = re.sub(r'[^\w\-_.]', '', nome_cliente)[:15]  # Limitar tamanho
#         timestamp = datetime.now().strftime('%d%m%Y_%H%M')
#         output_filename = f"doc_{nome_cliente}_{timestamp}.docx"
#         output_path = os.path.join(temp_dir, output_filename)
        
#         logger.info(f"Gerando documento para WhatsApp: {output_filename}")
        
#         try:
#             # Procurar template
#             possible_templates = [
#                 "template.docx",
#                 "modelo.docx",
#                 "templates/template.docx",
#                 "templates/modelo.docx"
#             ]
            
#             template_encontrado = None
#             for template_path in possible_templates:
#                 if os.path.exists(template_path):
#                     template_encontrado = template_path
#                     break
            
#             if template_encontrado:
#                 logger.info(f"Template encontrado: {template_encontrado}")
#                 services.preencher_modelo(template_encontrado, output_path, dados_extraidos)
#                 logger.info("Template preenchido com sucesso")
#             else:
#                 logger.info("Template não encontrado, criando documento padrão")
#                 services.criar_documento_fallback(dados_extraidos, output_path)
                
#         except Exception as e:
#             logger.error(f"Erro no preenchimento: {e}")
#             logger.info("Criando documento fallback...")
#             services.criar_documento_fallback(dados_extraidos, output_path)
        
#         # Verificar se arquivo foi criado
#         if not os.path.exists(output_path):
#             raise Exception("Documento não foi gerado")
        
#         file_size = os.path.getsize(output_path)
#         logger.info(f"Documento criado: {file_size} bytes")
        
#         # Verificar se o arquivo não está corrompido
#         if file_size < 1000:  # DOCX mínimo tem pelo menos 1KB
#             raise Exception("Arquivo gerado parece estar corrompido (muito pequeno)")
        
#         # Converter para base64 com validação
#         with open(output_path, "rb") as f:
#             docx_content = f.read()
            
#             # Validar se é um arquivo DOCX válido (inicia com PK)
#             if not docx_content.startswith(b'PK'):
#                 raise Exception("Arquivo gerado não é um DOCX válido")
            
#             # Gerar base64 limpo
#             base64_content = base64.b64encode(docx_content).decode('utf-8')
            
#             # Verificar se base64 foi gerado corretamente
#             if not base64_content or len(base64_content) < 100:
#                 raise Exception("Erro na codificação base64")
            
#             logger.info(f"Base64 gerado: {len(base64_content)} caracteres")
        
#         # Criar caption curta para WhatsApp
#         nome_curto = dados_extraidos.get('NOME', 'Cliente')[:30]
#         caption = f"📄 {nome_curto}\n📅 {dados_extraidos.get('DATA', 'N/A')} {dados_extraidos.get('HORA', 'N/A')}"
        
#         # Resposta otimizada para integração com Z-API
#         return {
#             "success": True,
#             "status": "document_ready",
#             "message": "Documento gerado com sucesso para WhatsApp",
#             "file": {
#                 "filename": output_filename,
#                 "base64": base64_content,
#                 "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 "caption": caption,
#                 "size": file_size
#             },
#             # Formato alternativo para diferentes APIs
#             "whatsapp_data": {
#                 "filename": output_filename,
#                 "base64": base64_content,
#                 "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 "caption": caption
#             },
#             "document_info": {
#                 "filename": output_filename,
#                 "file_size": file_size,
#                 "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 "base64_length": len(base64_content)
#             },
#             "dados_extraidos": dados_extraidos,
#             "timestamp": datetime.now().isoformat(),
#             "environment": "production"
#         }
        
#     except Exception as e:
#         logger.error(f"ERRO CRÍTICO: {e}")
#         return {
#             "success": False,
#             "status": "error",
#             "message": f"Erro na geração do documento: {str(e)}",
#             "timestamp": datetime.now().isoformat()
#         }
    
#     finally:
#         # Limpar arquivos temporários
#         try:
#             if temp_dir and os.path.exists(temp_dir):
#                 shutil.rmtree(temp_dir)
#                 logger.info("Arquivos temporários removidos")
#         except Exception as e:
#             logger.warning(f"Erro na limpeza: {e}")


# @app.post("/webhook/processar")
# async def webhook_processar(dados: dict):
#     """Endpoint específico para webhooks do N8N"""
#     logger.info("=== WEBHOOK N8N CLOUD ===")
#     logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
#     logger.info(f"Dados recebidos: {dados}")
    
#     try:
#         agora = datetime.now()
        
#         # Verificar se tem mensagem em texto
#         mensagem_texto = None
#         if "mensagem" in dados:
#             mensagem_texto = dados["mensagem"]
#         elif "message" in dados:
#             mensagem_texto = dados["message"]
#         elif "texto" in dados:
#             mensagem_texto = dados["texto"]
        
#         if mensagem_texto:
#             dados_extraidos = services.extrair_dados_da_mensagem(mensagem_texto)
#         else:
#             dados_extraidos = {
#                 "NOME": dados.get("nome") or dados.get("NOME") or "Não informado",
#                 "EMAIL": dados.get("email") or dados.get("EMAIL") or "Não informado",
#                 "CPF": dados.get("cpf") or dados.get("CPF") or "Não informado",
#                 "ENDERECO": dados.get("endereco") or dados.get("ENDERECO") or "Não informado",
#                 "CEP": dados.get("cep") or dados.get("CEP") or "Não informado",
#                 "TELEFONE": dados.get("telefone") or dados.get("TELEFONE") or "Não informado",
#                 "VALOR": dados.get("valor") or dados.get("VALOR") or "Não informado",
#                 "PARCELAS": dados.get("parcelas") or dados.get("PARCELAS") or "Não informado",
#                 "FORMA_PAGAMENTO": dados.get("forma_pagamento") or dados.get("FORMA_PAGAMENTO") or "Não informado",
#                 "PACIENTE": dados.get("nome") or dados.get("NOME") or "Não informado",
                
#                 "DATA": agora.strftime("%d/%m/%Y"),
#                 "HORA": agora.strftime("%H:%M:%S"),
#                 "DATA_HORA": agora.strftime("%d/%m/%Y %H:%M:%S"),
#                 "DATA_PROCESSAMENTO": agora.strftime("%d/%m/%Y %H:%M:%S"),
#                 "TIMESTAMP": agora.isoformat(),
#                 "ARQUIVO_FONTE": "Webhook N8N Cloud"
#             }
        
#         return {
#             "status": "success",
#             "message": "Dados processados com sucesso",
#             "dados": dados_extraidos,
#             "timestamp": agora.isoformat(),
#             "data_processamento": agora.strftime("%d/%m/%Y %H:%M:%S"),
#             "environment": "production"
#         }
        
#     except Exception as e:
#         logger.error(f"Erro no webhook: {e}")
#         return {
#             "status": "error",
#             "message": f"Erro no processamento: {str(e)}",
#             "timestamp": datetime.now().isoformat()
#         }


# @app.post("/gerar-documento-zapi")
# async def gerar_documento_zapi(request: MensagemRequest):
#     """Endpoint específico para Z-API com formato exato que ela espera"""
#     logger.info("=== GERAÇÃO DE DOCUMENTO PARA Z-API ===")
#     logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
#     temp_dir = tempfile.mkdtemp()
    
#     try:
#         # Extrair dados
#         dados_extraidos = services.extrair_dados_da_mensagem(request.mensagem)
#         logger.info("Dados extraídos para Z-API")
        
#         # Nome do arquivo simplificado
#         nome_cliente = re.sub(r'[^\w]', '', dados_extraidos.get("NOME", "cliente"))[:10]
#         timestamp = datetime.now().strftime('%d%m_%H%M')
#         filename = f"{nome_cliente}_{timestamp}.docx"
#         output_path = os.path.join(temp_dir, filename)
        
#         # Criar documento
#         try:
#             possible_templates = ["template.docx", "modelo.docx", "templates/template.docx"]
#             template_encontrado = None
#             for template_path in possible_templates:
#                 if os.path.exists(template_path):
#                     template_encontrado = template_path
#                     break
            
#             if template_encontrado:
#                 services.preencher_modelo(template_encontrado, output_path, dados_extraidos)
#             else:
#                 services.criar_documento_fallback(dados_extraidos, output_path)
                
#         except Exception as e:
#             logger.error(f"Erro: {e}")
#             services.criar_documento_fallback(dados_extraidos, output_path)
        
#         # Verificar arquivo
#         if not os.path.exists(output_path):
#             raise Exception("Documento não foi gerado")
        
#         file_size = os.path.getsize(output_path)
#         if file_size < 1000:
#             raise Exception("Arquivo muito pequeno - possível corrupção")
        
#         # Ler e validar arquivo
#         with open(output_path, "rb") as f:
#             file_bytes = f.read()
            
#         # Validar se é DOCX válido
#         if not file_bytes.startswith(b'PK'):
#             raise Exception("Arquivo não é um DOCX válido")
        
#         # Gerar base64 sem quebras de linha
#         base64_string = base64.b64encode(file_bytes).decode('ascii')
        
#         # Validar base64
#         if len(base64_string) < 1000:
#             raise Exception("Base64 muito pequeno")
        
#         # Testar se base64 pode ser decodificado
#         try:
#             base64.b64decode(base64_string)
#         except Exception:
#             raise Exception("Base64 inválido gerado")
        
#         logger.info(f"✅ Arquivo: {filename} ({file_size} bytes)")
#         logger.info(f"✅ Base64: {len(base64_string)} caracteres")
        
#         return {
#             "success": True,
#             "filename": filename,
#             "base64": base64_string,
#             "size": file_size,
#             "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             "dados": dados_extraidos,
#             "timestamp": datetime.now().isoformat()
#         }
        
#     except Exception as e:
#         logger.error(f"❌ ERRO: {e}")
#         return {
#             "success": False,
#             "error": str(e),
#             "timestamp": datetime.now().isoformat()
#         }
    
#     finally:
#         try:
#             if temp_dir and os.path.exists(temp_dir):
#                 shutil.rmtree(temp_dir)
#         except:
#             pass


# @app.post("/test-docx")
# async def test_docx():
#     """Endpoint para testar geração de DOCX simples"""
#     temp_dir = tempfile.mkdtemp()
    
#     try:
#         # Criar documento de teste
#         doc = Document()
#         doc.add_heading('Teste de Documento', 0)
#         doc.add_paragraph('Este é um teste de geração de DOCX.')
#         doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        
#         filename = f"teste_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx"
#         filepath = os.path.join(temp_dir, filename)
#         doc.save(filepath)
        
#         # Ler arquivo
#         with open(filepath, "rb") as f:
#             file_bytes = f.read()
        
#         # Gerar base64
#         base64_string = base64.b64encode(file_bytes).decode('ascii')
        
#         return {
#             "success": True,
#             "message": "Teste OK",
#             "filename": filename,
#             "size": len(file_bytes),
#             "base64_length": len(base64_string),
#             "base64_preview": base64_string[:100] + "...",
#             "is_valid_docx": file_bytes.startswith(b'PK'),
#             "timestamp": datetime.now().isoformat()
#         }
        
#     except Exception as e:
#         return {
#             "success": False,
#             "error": str(e),
#             "timestamp": datetime.now().isoformat()
#         }
    
#     finally:
#         try:
#             if temp_dir and os.path.exists(temp_dir):
#                 shutil.rmtree(temp_dir)
#         except:
#             pass


# # check if .docx is valid
# @app.post("/debug-template")
# async def debug_template():
#     """Endpoint para fazer debug de um template DOCX"""
#     try:
#         # Procurar template
#         possible_templates = [
#             "template.docx",
#             "modelo.docx", 
#             "templates/template.docx",
#             "templates/modelo.docx"
#         ]
        
#         template_encontrado = None
#         for template_path in possible_templates:
#             if os.path.exists(template_path):
#                 template_encontrado = template_path
#                 break
        
#         if not template_encontrado:
#             return {
#                 "success": False,
#                 "error": "Template não encontrado",
#                 "paths_testados": possible_templates,
#                 "timestamp": datetime.now().isoformat()
#             }
        
#         # Analisar template
#         doc = Document(template_encontrado)
        
#         # Dados de teste
#         dados_teste = {
#             "NOME": "João Silva",
#             "VALOR": "R$ 1.500,00",
#             "EMAIL": "joao@teste.com",
#             "CPF": "123.456.789-00"
#         }
        
#         # Verificar placeholders
#         placeholders_encontrados = services.verificar_placeholders_no_documento(doc, dados_teste)
        
#         # Análise de estrutura
#         analise = {
#             "template_path": template_encontrado,
#             "total_paragrafos": len(doc.paragraphs),
#             "total_tabelas": len(doc.tables),
#             "total_secoes": len(doc.sections),
#             "placeholders_encontrados": placeholders_encontrados,
#             "dados_teste": dados_teste
#         }
        
#         # Detalhes dos primeiros parágrafos
#         paragrafos_detalhes = []
#         for i, paragrafo in enumerate(doc.paragraphs[:10]):
#             texto_completo = ''.join(run.text for run in paragrafo.runs)
#             if texto_completo.strip():
#                 paragrafos_detalhes.append({
#                     "indice": i,
#                     "texto": texto_completo[:100],
#                     "tem_placeholder": '{{' in texto_completo,
#                     "num_runs": len(paragrafo.runs)
#                 })
        
#         analise["paragrafos_amostra"] = paragrafos_detalhes
        
#         return {
#             "success": True,
#             "analise": analise,
#             "timestamp": datetime.now().isoformat()
#         }
        
#     except Exception as e:
#         return {
#             "success": False,
#             "error": str(e),
#             "timestamp": datetime.now().isoformat()
#         }
