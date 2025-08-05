import re
from core import logger
from docx import Document
from datetime import datetime

def substituir_placeholders_robusto(paragrafos, dados):
    """
    Substitui placeholders de forma mais robusta, lidando com runs fragmentados
    """
    for paragrafo in paragrafos:
        if not paragrafo.runs:
            continue
            
        # Consolidar texto completo do parágrafo
        texto_completo = ''.join(run.text for run in paragrafo.runs)
        
        # Verificar se há placeholders no texto
        texto_modificado = texto_completo
        houve_substituicao = False
        
        for chave, valor in dados.items():
            placeholder = f'{{{{{chave}}}}}'
            valor_str = str(valor) if valor is not None else "Não informado"
            
            if placeholder in texto_modificado:
                texto_modificado = texto_modificado.replace(placeholder, valor_str)
                houve_substituicao = True
                logger.info(f"✅ Substituído: {placeholder} -> {valor_str}")
        
        # Se houve substituição, reconstruir o parágrafo
        if houve_substituicao:
            # Preservar formatação do primeiro run com conteúdo
            formatacao_base = None
            for run in paragrafo.runs:
                if run.text.strip():
                    formatacao_base = {
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline,
                        'color': run.font.color.rgb if run.font.color.rgb else None
                    }
                    break
            
            # Limpar todos os runs
            for run in paragrafo.runs:
                run.text = ""
            
            # Garantir que há pelo menos um run
            if not paragrafo.runs:
                paragrafo.add_run()
            
            # Aplicar texto modificado no primeiro run
            primeiro_run = paragrafo.runs[0]
            primeiro_run.text = texto_modificado
            
            # Aplicar formatação preservada
            if formatacao_base:
                try:
                    if formatacao_base['font_name']:
                        primeiro_run.font.name = formatacao_base['font_name']
                    if formatacao_base['font_size']:
                        primeiro_run.font.size = formatacao_base['font_size']
                    primeiro_run.font.bold = formatacao_base['bold'] or False
                    primeiro_run.font.italic = formatacao_base['italic'] or False
                    primeiro_run.font.underline = formatacao_base['underline'] or False
                    if formatacao_base['color']:
                        primeiro_run.font.color.rgb = formatacao_base['color']
                except Exception as e:
                    logger.warning(f"Erro ao aplicar formatação: {e}")

def verificar_placeholders_no_documento(doc, dados):
    """
    Verifica e lista todos os placeholders encontrados no documento
    """
    placeholders_encontrados = set()
    
    # Verificar parágrafos principais
    for paragrafo in doc.paragraphs:
        texto = ''.join(run.text for run in paragrafo.runs)
        matches = re.findall(r'\{\{([^}]+)\}\}', texto)
        placeholders_encontrados.update(matches)
    
    # Verificar tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto = ''.join(run.text for run in paragrafo.runs)
                    matches = re.findall(r'\{\{([^}]+)\}\}', texto)
                    placeholders_encontrados.update(matches)
    
    # Verificar cabeçalhos e rodapés
    for section in doc.sections:
        if section.header:
            for paragrafo in section.header.paragraphs:
                texto = ''.join(run.text for run in paragrafo.runs)
                matches = re.findall(r'\{\{([^}]+)\}\}', texto)
                placeholders_encontrados.update(matches)
        
        if section.footer:
            for paragrafo in section.footer.paragraphs:
                texto = ''.join(run.text for run in paragrafo.runs)
                matches = re.findall(r'\{\{([^}]+)\}\}', texto)
                placeholders_encontrados.update(matches)
    
    logger.info(f"📋 Placeholders encontrados no documento: {list(placeholders_encontrados)}")
    logger.info(f"📊 Dados disponíveis para substituição: {list(dados.keys())}")
    
    # Verificar quais placeholders não têm dados correspondentes
    sem_dados = [p for p in placeholders_encontrados if p not in dados]
    if sem_dados:
        logger.warning(f"⚠️ Placeholders sem dados correspondentes: {sem_dados}")
    
    return list(placeholders_encontrados)

def debug_documento_runs(doc, limite_paragrafos=5):
    """
    Função para debug - mostra como os runs estão organizados no documento
    """
    logger.info("🔍 DEBUG: Analisando estrutura de runs do documento")
    
    for i, paragrafo in enumerate(doc.paragraphs[:limite_paragrafos]):
        if not paragrafo.runs:
            continue
            
        texto_completo = ''.join(run.text for run in paragrafo.runs)
        if '{{' in texto_completo:
            logger.info(f"📍 Parágrafo {i}: '{texto_completo[:100]}...'")
            logger.info(f"   Número de runs: {len(paragrafo.runs)}")
            
            for j, run in enumerate(paragrafo.runs):
                if run.text:
                    logger.info(f"   Run {j}: '{run.text}'")

def preencher_modelo(caminho_modelo, caminho_saida, dados):
    """Preenche um modelo DOCX com os dados fornecidos - VERSÃO CORRIGIDA"""
    try:
        logger.info(f"📖 Abrindo modelo: {caminho_modelo}")
        doc = Document(caminho_modelo)
        
        # Preparar dados - garantir que todos os valores sejam strings
        dados_limpos = {}
        for chave, valor in dados.items():
            if valor is None or valor == "":
                dados_limpos[chave] = "Não informado"
            else:
                dados_limpos[chave] = str(valor).strip()
        
        logger.info(f"📋 Dados preparados para substituição:")
        for chave, valor in dados_limpos.items():
            logger.info(f"   {chave}: {valor}")
        
        # Debug: verificar estrutura do documento se necessário
        debug_documento_runs(doc, limite_paragrafos=3)
        
        # Verificar placeholders antes da substituição
        placeholders_iniciais = verificar_placeholders_no_documento(doc, dados_limpos)
        
        # Processar parágrafos principais
        logger.info("📄 Processando parágrafos principais...")
        substituir_placeholders_robusto(doc.paragraphs, dados_limpos)
        
        # Processar tabelas
        logger.info("📊 Processando tabelas...")
        for i, tabela in enumerate(doc.tables):
            for j, linha in enumerate(tabela.rows):
                for k, celula in enumerate(linha.cells):
                    substituir_placeholders_robusto(celula.paragraphs, dados_limpos)
        
        # Processar cabeçalhos e rodapés
        logger.info("📑 Processando cabeçalhos e rodapés...")
        for i, section in enumerate(doc.sections):
            if section.header:
                substituir_placeholders_robusto(section.header.paragraphs, dados_limpos)
            if section.footer:
                substituir_placeholders_robusto(section.footer.paragraphs, dados_limpos)
        
        # Salvar documento
        logger.info(f"💾 Salvando documento em: {caminho_saida}")
        doc.save(caminho_saida)
        
        # Verificação final
        doc_verificacao = Document(caminho_saida)
        placeholders_restantes = verificar_placeholders_no_documento(doc_verificacao, dados_limpos)
        
        if placeholders_restantes:
            logger.warning(f"⚠️ ATENÇÃO: Ainda existem placeholders não substituídos: {placeholders_restantes}")
            
            # Tentar substituição adicional mais agressiva
            logger.info("🔧 Tentando substituição adicional...")
            for paragrafo in doc_verificacao.paragraphs:
                texto_original = paragrafo.text
                if '{{' in texto_original:
                    # Substituição direta no texto do parágrafo
                    novo_texto = texto_original
                    for chave, valor in dados_limpos.items():
                        placeholder = f'{{{{{chave}}}}}'
                        if placeholder in novo_texto:
                            novo_texto = novo_texto.replace(placeholder, valor)
                    
                    if novo_texto != texto_original:
                        # Limpar runs e recriar
                        for run in paragrafo.runs:
                            run.text = ""
                        if paragrafo.runs:
                            paragrafo.runs[0].text = novo_texto
                        else:
                            paragrafo.add_run(novo_texto)
                        
                        logger.info(f"🔧 Correção aplicada: '{texto_original[:50]}...' -> '{novo_texto[:50]}...'")
            
            # Processar tabelas na verificação final
            for tabela in doc_verificacao.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            texto_original = paragrafo.text
                            if '{{' in texto_original:
                                novo_texto = texto_original
                                for chave, valor in dados_limpos.items():
                                    placeholder = f'{{{{{chave}}}}}'
                                    if placeholder in novo_texto:
                                        novo_texto = novo_texto.replace(placeholder, valor)
                                
                                if novo_texto != texto_original:
                                    for run in paragrafo.runs:
                                        run.text = ""
                                    if paragrafo.runs:
                                        paragrafo.runs[0].text = novo_texto
                                    else:
                                        paragrafo.add_run(novo_texto)
                                    
                                    logger.info(f"🔧 Correção em tabela: '{texto_original[:30]}...' -> '{novo_texto[:30]}...'")
            
            # Salvar novamente após correções
            doc_verificacao.save(caminho_saida)
            
            # Verificação final final
            doc_final = Document(caminho_saida)
            placeholders_finais = verificar_placeholders_no_documento(doc_final, dados_limpos)
            
            if placeholders_finais:
                logger.error(f"❌ AINDA RESTAM placeholders não substituídos: {placeholders_finais}")
            else:
                logger.info("✅ Correção bem-sucedida! Todos os placeholders foram substituídos.")
        else:
            logger.info("✅ Todos os placeholders foram substituídos com sucesso!")
        
        logger.info("✅ Processamento concluído!")
        return True
        
    except Exception as e:
        logger.error(f"❌ Erro ao preencher modelo: {str(e)}")
        raise Exception(f"Erro ao preencher modelo: {str(e)}")

def extrair_dados_da_mensagem(mensagem: str) -> dict:
    """Extrai os dados da mensagem com validação aprimorada"""
    dados = {}
    
    # Log da mensagem recebida para debug
    logger.info(f"📨 Mensagem recebida para extração:")
    logger.info(f"   Tamanho: {len(mensagem)} caracteres")
    logger.info(f"   Prévia: {mensagem[:200]}...")
    
    padroes = {
        "NOME": [
            r"Nome:\s*(.+?)(?=\n|$)", 
            r"nome:\s*(.+?)(?=\n|$)",
            r"NOME:\s*(.+?)(?=\n|$)"
        ],
        "EMAIL": [
            r"Email:\s*(.+?)(?=\n|$)", 
            r"email:\s*(.+?)(?=\n|$)", 
            r"E-mail:\s*(.+?)(?=\n|$)",
            r"EMAIL:\s*(.+?)(?=\n|$)"
        ],
        "CPF": [
            r"CPF:\s*(.+?)(?=\n|$)", 
            r"cpf:\s*(.+?)(?=\n|$)"
        ],
        "ENDERECO": [
            r"Endereço:\s*(.+?)(?=\n|$)", 
            r"endereco:\s*(.+?)(?=\n|$)", 
            r"Endereco:\s*(.+?)(?=\n|$)",
            r"ENDERECO:\s*(.+?)(?=\n|$)"
        ],
        "CEP": [
            r"CEP:\s*(.+?)(?=\n|$)", 
            r"cep:\s*(.+?)(?=\n|$)"
        ],
        "TELEFONE": [
            r"Telefone:\s*(.+?)(?=\n|$)", 
            r"telefone:\s*(.+?)(?=\n|$)", 
            r"Fone:\s*(.+?)(?=\n|$)",
            r"TELEFONE:\s*(.+?)(?=\n|$)"
        ],
        "VALOR": [
            r"Valor:\s*(.+?)(?=\n|$)", 
            r"valor:\s*(.+?)(?=\n|$)",
            r"VALOR:\s*(.+?)(?=\n|$)"
        ],
        "PARCELAS": [
            r"Quantidade de Parcelas:\s*(.+?)(?=\n|$)", 
            r"quantidade de parcelas:\s*(.+?)(?=\n|$)",
            r"Parcelas:\s*(.+?)(?=\n|$)",
            r"parcelas:\s*(.+?)(?=\n|$)",
            r"PARCELAS:\s*(.+?)(?=\n|$)"
        ],
        "FORMA_PAGAMENTO": [
            r"Forma de pagamento:\s*(.+?)(?=\n|$)", 
            r"forma de pagamento:\s*(.+?)(?=\n|$)",
            r"Pagamento:\s*(.+?)(?=\n|$)",
            r"pagamento:\s*(.+?)(?=\n|$)",
            r"FORMA_PAGAMENTO:\s*(.+?)(?=\n|$)"
        ]
    }
    
    # Extrair dados usando os padrões
    for campo, padroes_campo in padroes.items():
        valor_encontrado = None
        for padrao in padroes_campo:
            match = re.search(padrao, mensagem, re.IGNORECASE | re.MULTILINE)
            if match:
                valor_encontrado = match.group(1).strip()
                logger.info(f"✅ {campo}: {valor_encontrado}")
                break
        
        dados[campo] = valor_encontrado if valor_encontrado else "Não informado"
        
        if not valor_encontrado:
            logger.info(f"⚠️ {campo}: Não encontrado")
    
    # Adicionar campos de data/hora automaticamente
    agora = datetime.now()
    dados["DATA"] = agora.strftime("%d/%m/%Y")
    dados["HORA"] = agora.strftime("%H:%M:%S")
    dados["DATA_HORA"] = agora.strftime("%d/%m/%Y %H:%M:%S")
    dados["DATA_PROCESSAMENTO"] = agora.strftime("%d/%m/%Y %H:%M:%S")
    dados["TIMESTAMP"] = agora.isoformat()
    
    # Campos derivados
    dados["PACIENTE"] = dados["NOME"]
    dados["ARQUIVO_FONTE"] = "API N8N Cloud"
    
    logger.info(f"📊 Resumo da extração:")
    logger.info(f"   Total de campos extraídos: {len([v for v in dados.values() if v != 'Não informado'])}")
    logger.info(f"   Campos sem valor: {len([v for v in dados.values() if v == 'Não informado'])}")
    
    return dados

def criar_documento_fallback(dados: dict, output_path: str) -> None:
    """Cria um documento DOCX simples com os dados extraídos"""
    doc = Document()
    
    # Cabeçalho
    doc.add_heading('Dados do Cliente', 0)
    doc.add_paragraph(f'Processado em: {dados.get("DATA_HORA", "N/A")}')
    doc.add_paragraph('---')
    
    # Seção de informações pessoais
    doc.add_heading('Informações Pessoais', level=1)
    doc.add_paragraph(f'Nome: {dados.get("NOME", "Não informado")}')
    doc.add_paragraph(f'Email: {dados.get("EMAIL", "Não informado")}')
    doc.add_paragraph(f'CPF: {dados.get("CPF", "Não informado")}')
    doc.add_paragraph(f'Telefone: {dados.get("TELEFONE", "Não informado")}')
    
    # Seção de endereço
    doc.add_heading('Endereço', level=1)
    doc.add_paragraph(f'Endereço: {dados.get("ENDERECO", "Não informado")}')
    doc.add_paragraph(f'CEP: {dados.get("CEP", "Não informado")}')
    
    # Seção financeira
    doc.add_heading('Informações Financeiras', level=1)
    doc.add_paragraph(f'Valor: {dados.get("VALOR", "Não informado")}')
    doc.add_paragraph(f'Quantidade de Parcelas: {dados.get("PARCELAS", "Não informado")}')
    doc.add_paragraph(f'Forma de Pagamento: {dados.get("FORMA_PAGAMENTO", "Não informado")}')
    
    # Adicionar data/hora
    doc.add_heading('Informações do Processamento', level=1)
    doc.add_paragraph(f'Data: {dados.get("DATA", "N/A")}')
    doc.add_paragraph(f'Hora: {dados.get("HORA", "N/A")}')
    
    doc.save(output_path)
