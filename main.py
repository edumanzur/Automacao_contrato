from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel
import os
import tempfile
import shutil
import re
from docx import Document
from datetime import datetime
from typing import Optional
import logging
import base64
import mimetypes

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="API Processamento de Mensagens N8N + WhatsApp",
    description="API para processar mensagens do N8N e gerar documentos DOCX para WhatsApp",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS para permitir acesso do N8N
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Em produção, especifique as origens
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Modelos Pydantic
class MensagemRequest(BaseModel):
    mensagem: str
    webhook_id: Optional[str] = None
    origem: Optional[str] = "n8n"
    formato_resposta: Optional[str] = "binary"  # binary, base64, json

class MensagemResponse(BaseModel):
    sucesso: bool
    mensagem: str
    dados_extraidos: dict
    arquivo_gerado: Optional[str] = None

class DocumentoResponse(BaseModel):
    success: bool
    message: str
    filename: str
    file_size: int
    mime_type: str
    base64_content: Optional[str] = None
    download_url: Optional[str] = None
    dados_extraidos: dict
    timestamp: str

def substituir_em_runs_preservando_formatacao(paragrafos, dados):
    """Substitui placeholders nos runs preservando TODA a formatação original"""
    for paragrafo in paragrafos:
        # Primeiro, identificar todos os placeholders no parágrafo completo
        texto_completo = ""
        runs_info = []
        
        for i, run in enumerate(paragrafo.runs):
            runs_info.append({
                'index': i,
                'text': run.text,
                'start_pos': len(texto_completo),
                'end_pos': len(texto_completo) + len(run.text),
                'font': {
                    'name': run.font.name,
                    'size': run.font.size,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': run.font.color.rgb if run.font.color.rgb else None
                }
            })
            texto_completo += run.text
        
        # Verificar se há placeholders para substituir
        texto_modificado = texto_completo
        substituicoes_feitas = False
        mapa_substituicoes = []
        
        for chave, valor in dados.items():
            placeholder = f'{{{{{chave}}}}}'
            if placeholder in texto_modificado:
                valor_str = str(valor) if valor is not None else ""
                
                # Registrar posições das substituições
                start_pos = 0
                while True:
                    pos = texto_modificado.find(placeholder, start_pos)
                    if pos == -1:
                        break
                    
                    mapa_substituicoes.append({
                        'placeholder': placeholder,
                        'valor': valor_str,
                        'pos_inicio': pos,
                        'pos_fim': pos + len(placeholder),
                        'tamanho_original': len(placeholder),
                        'tamanho_novo': len(valor_str)
                    })
                    
                    start_pos = pos + 1
                
                texto_modificado = texto_modificado.replace(placeholder, valor_str)
                substituicoes_feitas = True
                logger.info(f"Substituído: {placeholder} -> {valor_str}")
        
        # Se houve substituições, reconstruir os runs preservando formatação
        if substituicoes_feitas:
            # Limpar runs existentes
            for run in paragrafo.runs:
                run.text = ""
            
            # Se não há runs, criar um
            if not paragrafo.runs:
                paragrafo.add_run()
            
            # Método simplificado: aplicar todo o texto modificado no primeiro run
            # e copiar a formatação do run original que continha a maior parte do placeholder
            primeiro_run = paragrafo.runs[0]
            primeiro_run.text = texto_modificado
            
            # Tentar preservar a formatação do primeiro run não vazio original
            for run_info in runs_info:
                if run_info['text'].strip():  # Primeiro run com texto
                    if run_info['font']['name']:
                        primeiro_run.font.name = run_info['font']['name']
                    if run_info['font']['size']:
                        primeiro_run.font.size = run_info['font']['size']
                    primeiro_run.font.bold = run_info['font']['bold']
                    primeiro_run.font.italic = run_info['font']['italic']
                    primeiro_run.font.underline = run_info['font']['underline']
                    if run_info['font']['color']:
                        primeiro_run.font.color.rgb = run_info['font']['color']
                    break

def substituir_placeholder_inteligente(paragrafos, dados):
    """Versão melhorada que preserva formatação por meio de substituição inteligente"""
    for paragrafo in paragrafos:
        # Construir mapa de posições dos runs
        posicoes_runs = []
        texto_completo = ""
        
        for i, run in enumerate(paragrafo.runs):
            inicio = len(texto_completo)
            fim = inicio + len(run.text)
            posicoes_runs.append({
                'run': run,
                'inicio': inicio,
                'fim': fim,
                'texto_original': run.text
            })
            texto_completo += run.text
        
        # Verificar substituições necessárias
        substituicoes_realizadas = False
        
        for chave, valor in dados.items():
            placeholder = f'{{{{{chave}}}}}'
            valor_str = str(valor) if valor is not None else ""
            
            if placeholder in texto_completo:
                # Encontrar qual(is) run(s) contém(êm) o placeholder
                pos_placeholder = texto_completo.find(placeholder)
                
                if pos_placeholder != -1:
                    # Identificar o run que contém o início do placeholder
                    run_alvo = None
                    for pos_info in posicoes_runs:
                        if pos_info['inicio'] <= pos_placeholder < pos_info['fim']:
                            run_alvo = pos_info['run']
                            break
                    
                    if run_alvo:
                        # Fazer a substituição preservando a formatação do run
                        novo_texto = run_alvo.text.replace(placeholder, valor_str)
                        run_alvo.text = novo_texto
                        substituicoes_realizadas = True
                        logger.info(f"Substituído {placeholder} -> {valor_str} (formatação preservada)")

def preencher_modelo(caminho_modelo, caminho_saida, dados):
    """Preenche um modelo DOCX com os dados fornecidos"""
    try:
        logger.info(f"Abrindo modelo: {caminho_modelo}")
        doc = Document(caminho_modelo)
        
        dados_limpos = {}
        for chave, valor in dados.items():
            if valor is None or valor == "":
                dados_limpos[chave] = "Não informado"
            else:
                dados_limpos[chave] = str(valor)
        
        logger.info("Processando parágrafos principais...")
        substituir_placeholder_inteligente(doc.paragraphs, dados_limpos)
        
        logger.info("Processando tabelas...")
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    substituir_placeholder_inteligente(celula.paragraphs, dados_limpos)
        
        logger.info("Processando cabeçalhos e rodapés...")
        for section in doc.sections:
            if section.header:
                substituir_placeholder_inteligente(section.header.paragraphs, dados_limpos)
            if section.footer:
                substituir_placeholder_inteligente(section.footer.paragraphs, dados_limpos)
        
        logger.info(f"Salvando documento em: {caminho_saida}")
        doc.save(caminho_saida)
        logger.info("Arquivo gerado com sucesso!")
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao preencher modelo: {str(e)}")
        raise Exception(f"Erro ao preencher modelo: {str(e)}")

def extrair_dados_da_mensagem(mensagem: str) -> dict:
    """Extrai os dados da mensagem com a estrutura especificada"""
    dados = {}
    
    padroes = {
        "NOME": [r"Nome:\s*(.+?)(?=\n|$)", r"nome:\s*(.+?)(?=\n|$)"],
        "EMAIL": [r"Email:\s*(.+?)(?=\n|$)", r"email:\s*(.+?)(?=\n|$)", r"E-mail:\s*(.+?)(?=\n|$)"],
        "CPF": [r"CPF:\s*(.+?)(?=\n|$)", r"cpf:\s*(.+?)(?=\n|$)"],
        "ENDERECO": [r"Endereço:\s*(.+?)(?=\n|$)", r"endereco:\s*(.+?)(?=\n|$)", r"Endereco:\s*(.+?)(?=\n|$)"],
        "CEP": [r"CEP:\s*(.+?)(?=\n|$)", r"cep:\s*(.+?)(?=\n|$)"],
        "TELEFONE": [r"Telefone:\s*(.+?)(?=\n|$)", r"telefone:\s*(.+?)(?=\n|$)", r"Fone:\s*(.+?)(?=\n|$)"],
        "VALOR": [r"Valor:\s*(.+?)(?=\n|$)", r"valor:\s*(.+?)(?=\n|$)"],
        "PARCELAS": [
            r"Quantidade de Parcelas:\s*(.+?)(?=\n|$)", 
            r"quantidade de parcelas:\s*(.+?)(?=\n|$)",
            r"Parcelas:\s*(.+?)(?=\n|$)",
            r"parcelas:\s*(.+?)(?=\n|$)"
        ],
        "FORMA_PAGAMENTO": [
            r"Forma de pagamento:\s*(.+?)(?=\n|$)", 
            r"forma de pagamento:\s*(.+?)(?=\n|$)",
            r"Pagamento:\s*(.+?)(?=\n|$)",
            r"pagamento:\s*(.+?)(?=\n|$)"
        ]
    }
    
    for campo, padroes_campo in padroes.items():
        valor_encontrado = None
        for padrao in padroes_campo:
            match = re.search(padrao, mensagem, re.IGNORECASE | re.MULTILINE)
            if match:
                valor_encontrado = match.group(1).strip()
                break
        
        dados[campo] = valor_encontrado if valor_encontrado else "Não informado"
    
    # Adicionar campos de data/hora automaticamente
    agora = datetime.now()
    dados["DATA"] = agora.strftime("%d/%m/%Y")
    dados["HORA"] = agora.strftime("%H:%M:%S")
    dados["DATA_HORA"] = agora.strftime("%d/%m/%Y %H:%M:%S")
    dados["DATA_PROCESSAMENTO"] = agora.strftime("%d/%m/%Y %H:%M:%S")
    dados["TIMESTAMP"] = agora.isoformat()
    
    dados["PACIENTE"] = dados["NOME"]
    dados["ARQUIVO_FONTE"] = "API N8N Cloud"
    
    return dados

def criar_documento_fallback(dados: dict, output_path: str) -> None:
    """Cria um documento DOCX simples com os dados extraídos"""
    doc = Document()
    
    # Cabeçalho
    doc.add_heading('Dados do Cliente', 0)
    doc.add_paragraph(f'Processado em: {dados.get("DATA_HORA", "N/A")}')
    doc.add_paragraph('---')
    
    # Seção de informações pessoais
    doc.add_heading('Informações Pessoais', level=1)
    doc.add_paragraph(f'Nome: {dados.get("NOME", "Não informado")}')
    doc.add_paragraph(f'Email: {dados.get("EMAIL", "Não informado")}')
    doc.add_paragraph(f'CPF: {dados.get("CPF", "Não informado")}')
    doc.add_paragraph(f'Telefone: {dados.get("TELEFONE", "Não informado")}')
    
    # Seção de endereço
    doc.add_heading('Endereço', level=1)
    doc.add_paragraph(f'Endereço: {dados.get("ENDERECO", "Não informado")}')
    doc.add_paragraph(f'CEP: {dados.get("CEP", "Não informado")}')
    
    # Seção financeira
    doc.add_heading('Informações Financeiras', level=1)
    doc.add_paragraph(f'Valor: {dados.get("VALOR", "Não informado")}')
    doc.add_paragraph(f'Quantidade de Parcelas: {dados.get("PARCELAS", "Não informado")}')
    doc.add_paragraph(f'Forma de Pagamento: {dados.get("FORMA_PAGAMENTO", "Não informado")}')
    
    # Adicionar data/hora
    doc.add_heading('Informações do Processamento', level=1)
    doc.add_paragraph(f'Data: {dados.get("DATA", "N/A")}')
    doc.add_paragraph(f'Hora: {dados.get("HORA", "N/A")}')
    
    doc.save(output_path)

@app.get("/")
async def root():
    return {
        "message": "API Processamento de Mensagens N8N + WhatsApp - Cloud Version",
        "version": "2.0.1",
        "status": "online",
        "environment": "production",
        "timestamp": datetime.now().isoformat(),
        "endpoints": {
            "gerar_documento": "POST /gerar-documento (retorna binário)",
            "gerar_documento_base64": "POST /gerar-documento-base64 (retorna JSON com base64)",
            "gerar_documento_whatsapp": "POST /gerar-documento-whatsapp (otimizado para Z-API)",
            "webhook": "POST /webhook/processar",
            "health": "GET /health"
        }
    }

@app.get("/health")
async def health_check():
    """Health check para o N8N verificar se a API está funcionando"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "processamento-mensagens-whatsapp-cloud",
        "environment": "production",
        "data_atual": datetime.now().strftime("%d/%m/%Y"),
        "hora_atual": datetime.now().strftime("%H:%M:%S")
    }

@app.post("/gerar-documento")
async def gerar_documento(request: MensagemRequest):
    """Endpoint para processar mensagem E gerar documento DOCX (retorna binário)"""
    logger.info("=== GERAÇÃO DE DOCUMENTO N8N CLOUD (BINÁRIO) ===")
    logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    temp_dir = tempfile.mkdtemp()
    output_path = None
    
    try:
        # Extrair dados
        dados_extraidos = extrair_dados_da_mensagem(request.mensagem)
        logger.info("Dados extraídos para documento")
        
        # Definir nome do arquivo
        nome_cliente = dados_extraidos.get("NOME", "cliente").replace(" ", "_")
        nome_cliente = re.sub(r'[^\w\-_.]', '', nome_cliente)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"documento_{nome_cliente}_{timestamp}.docx"
        output_path = os.path.join(temp_dir, output_filename)
        
        logger.info(f"Gerando documento: {output_filename}")
        
        try:
            # Procurar template
            possible_templates = [
                "template.docx",
                "modelo.docx",
                "templates/template.docx",
                "templates/modelo.docx"
            ]
            
            template_encontrado = None
            for template_path in possible_templates:
                if os.path.exists(template_path):
                    template_encontrado = template_path
                    break
            
            if template_encontrado:
                logger.info(f"Template encontrado: {template_encontrado}")
                preencher_modelo(template_encontrado, output_path, dados_extraidos)
                logger.info("Template preenchido com sucesso")
            else:
                logger.info("Template não encontrado, criando documento padrão")
                criar_documento_fallback(dados_extraidos, output_path)
                
        except Exception as e:
            logger.error(f"Erro no preenchimento: {e}")
            logger.info("Criando documento fallback...")
            criar_documento_fallback(dados_extraidos, output_path)
        
        # Verificar se arquivo foi criado
        if not os.path.exists(output_path):
            raise Exception("Documento não foi gerado")
        
        file_size = os.path.getsize(output_path)
        logger.info(f"Documento criado: {file_size} bytes")
        
        # Ler arquivo
        with open(output_path, "rb") as f:
            docx_content = f.read()
        
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={output_filename}",
                "Content-Length": str(len(docx_content)),
                "Cache-Control": "no-cache",
                "X-Filename": output_filename,
                "X-File-Size": str(file_size)
            }
        )
        
    except Exception as e:
        logger.error(f"ERRO CRÍTICO: {e}")
        raise HTTPException(status_code=500, detail=f"Erro na geração do documento: {str(e)}")
    
    finally:
        # Limpar arquivos temporários
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                logger.info("Arquivos temporários removidos")
        except Exception as e:
            logger.warning(f"Erro na limpeza: {e}")

@app.post("/gerar-documento-base64", response_model=DocumentoResponse)
async def gerar_documento_base64(request: MensagemRequest):
    """Endpoint que retorna o documento em base64 (ideal para integração com APIs)"""
    logger.info("=== GERAÇÃO DE DOCUMENTO N8N CLOUD (BASE64) ===")
    logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    temp_dir = tempfile.mkdtemp()
    output_path = None
    
    try:
        # Extrair dados
        dados_extraidos = extrair_dados_da_mensagem(request.mensagem)
        logger.info("Dados extraídos para documento")
        
        # Definir nome do arquivo
        nome_cliente = dados_extraidos.get("NOME", "cliente").replace(" ", "_")
        nome_cliente = re.sub(r'[^\w\-_.]', '', nome_cliente)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"documento_{nome_cliente}_{timestamp}.docx"
        output_path = os.path.join(temp_dir, output_filename)
        
        logger.info(f"Gerando documento: {output_filename}")
        
        try:
            # Procurar template
            possible_templates = [
                "template.docx",
                "modelo.docx",
                "templates/template.docx",
                "templates/modelo.docx"
            ]
            
            template_encontrado = None
            for template_path in possible_templates:
                if os.path.exists(template_path):
                    template_encontrado = template_path
                    break
            
            if template_encontrado:
                logger.info(f"Template encontrado: {template_encontrado}")
                preencher_modelo(template_encontrado, output_path, dados_extraidos)
                logger.info("Template preenchido com sucesso")
            else:
                logger.info("Template não encontrado, criando documento padrão")
                criar_documento_fallback(dados_extraidos, output_path)
                
        except Exception as e:
            logger.error(f"Erro no preenchimento: {e}")
            logger.info("Criando documento fallback...")
            criar_documento_fallback(dados_extraidos, output_path)
        
        # Verificar se arquivo foi criado
        if not os.path.exists(output_path):
            raise Exception("Documento não foi gerado")
        
        file_size = os.path.getsize(output_path)
        logger.info(f"Documento criado: {file_size} bytes")
        
        # Converter para base64
        with open(output_path, "rb") as f:
            docx_content = f.read()
            base64_content = base64.b64encode(docx_content).decode('utf-8')
        
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return DocumentoResponse(
            success=True,
            message="Documento gerado com sucesso",
            filename=output_filename,
            file_size=file_size,
            mime_type=mime_type,
            base64_content=base64_content,
            dados_extraidos=dados_extraidos,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        logger.error(f"ERRO CRÍTICO: {e}")
        raise HTTPException(status_code=500, detail=f"Erro na geração do documento: {str(e)}")
    
    finally:
        # Limpar arquivos temporários
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                logger.info("Arquivos temporários removidos")
        except Exception as e:
            logger.warning(f"Erro na limpeza: {e}")

@app.post("/gerar-documento-whatsapp")
async def gerar_documento_whatsapp(request: MensagemRequest):
    """Endpoint otimizado para envio via WhatsApp usando Z-API"""
    logger.info("=== GERAÇÃO DE DOCUMENTO PARA WHATSAPP ===")
    logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    temp_dir = tempfile.mkdtemp()
    output_path = None
    
    try:
        # Extrair dados
        dados_extraidos = extrair_dados_da_mensagem(request.mensagem)
        logger.info("Dados extraídos para documento WhatsApp")
        
        # Definir nome do arquivo (mais curto para WhatsApp)
        nome_cliente = dados_extraidos.get("NOME", "cliente").replace(" ", "_")
        nome_cliente = re.sub(r'[^\w\-_.]', '', nome_cliente)[:15]  # Limitar tamanho
        timestamp = datetime.now().strftime('%d%m%Y_%H%M')
        output_filename = f"doc_{nome_cliente}_{timestamp}.docx"
        output_path = os.path.join(temp_dir, output_filename)
        
        logger.info(f"Gerando documento para WhatsApp: {output_filename}")
        
        try:
            # Procurar template
            possible_templates = [
                "template.docx",
                "modelo.docx",
                "templates/template.docx",
                "templates/modelo.docx"
            ]
            
            template_encontrado = None
            for template_path in possible_templates:
                if os.path.exists(template_path):
                    template_encontrado = template_path
                    break
            
            if template_encontrado:
                logger.info(f"Template encontrado: {template_encontrado}")
                preencher_modelo(template_encontrado, output_path, dados_extraidos)
                logger.info("Template preenchido com sucesso")
            else:
                logger.info("Template não encontrado, criando documento padrão")
                criar_documento_fallback(dados_extraidos, output_path)
                
        except Exception as e:
            logger.error(f"Erro no preenchimento: {e}")
            logger.info("Criando documento fallback...")
            criar_documento_fallback(dados_extraidos, output_path)
        
        # Verificar se arquivo foi criado
        if not os.path.exists(output_path):
            raise Exception("Documento não foi gerado")
        
        file_size = os.path.getsize(output_path)
        logger.info(f"Documento criado: {file_size} bytes")
        
        # Verificar se o arquivo não está corrompido
        if file_size < 1000:  # DOCX mínimo tem pelo menos 1KB
            raise Exception("Arquivo gerado parece estar corrompido (muito pequeno)")
        
        # Converter para base64 com validação
        with open(output_path, "rb") as f:
            docx_content = f.read()
            
            # Validar se é um arquivo DOCX válido (inicia com PK)
            if not docx_content.startswith(b'PK'):
                raise Exception("Arquivo gerado não é um DOCX válido")
            
            # Gerar base64 limpo
            base64_content = base64.b64encode(docx_content).decode('utf-8')
            
            # Verificar se base64 foi gerado corretamente
            if not base64_content or len(base64_content) < 100:
                raise Exception("Erro na codificação base64")
            
            logger.info(f"Base64 gerado: {len(base64_content)} caracteres")
        
        # Criar caption curta para WhatsApp
        nome_curto = dados_extraidos.get('NOME', 'Cliente')[:30]
        caption = f"📄 {nome_curto}\n📅 {dados_extraidos.get('DATA', 'N/A')} {dados_extraidos.get('HORA', 'N/A')}"
        
        # Resposta otimizada para integração com Z-API
        return {
            "success": True,
            "status": "document_ready",
            "message": "Documento gerado com sucesso para WhatsApp",
            "file": {
                "filename": output_filename,
                "base64": base64_content,
                "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "caption": caption,
                "size": file_size
            },
            # Formato alternativo para diferentes APIs
            "whatsapp_data": {
                "filename": output_filename,
                "base64": base64_content,
                "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "caption": caption
            },
            "document_info": {
                "filename": output_filename,
                "file_size": file_size,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "base64_length": len(base64_content)
            },
            "dados_extraidos": dados_extraidos,
            "timestamp": datetime.now().isoformat(),
            "environment": "production"
        }
        
    except Exception as e:
        logger.error(f"ERRO CRÍTICO: {e}")
        return {
            "success": False,
            "status": "error",
            "message": f"Erro na geração do documento: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }
    
    finally:
        # Limpar arquivos temporários
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                logger.info("Arquivos temporários removidos")
        except Exception as e:
            logger.warning(f"Erro na limpeza: {e}")

@app.post("/webhook/processar")
async def webhook_processar(dados: dict):
    """Endpoint específico para webhooks do N8N"""
    logger.info("=== WEBHOOK N8N CLOUD ===")
    logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    logger.info(f"Dados recebidos: {dados}")
    
    try:
        agora = datetime.now()
        
        # Verificar se tem mensagem em texto
        mensagem_texto = None
        if "mensagem" in dados:
            mensagem_texto = dados["mensagem"]
        elif "message" in dados:
            mensagem_texto = dados["message"]
        elif "texto" in dados:
            mensagem_texto = dados["texto"]
        
        if mensagem_texto:
            dados_extraidos = extrair_dados_da_mensagem(mensagem_texto)
        else:
            dados_extraidos = {
                "NOME": dados.get("nome") or dados.get("NOME") or "Não informado",
                "EMAIL": dados.get("email") or dados.get("EMAIL") or "Não informado",
                "CPF": dados.get("cpf") or dados.get("CPF") or "Não informado",
                "ENDERECO": dados.get("endereco") or dados.get("ENDERECO") or "Não informado",
                "CEP": dados.get("cep") or dados.get("CEP") or "Não informado",
                "TELEFONE": dados.get("telefone") or dados.get("TELEFONE") or "Não informado",
                "VALOR": dados.get("valor") or dados.get("VALOR") or "Não informado",
                "PARCELAS": dados.get("parcelas") or dados.get("PARCELAS") or "Não informado",
                "FORMA_PAGAMENTO": dados.get("forma_pagamento") or dados.get("FORMA_PAGAMENTO") or "Não informado",
                "PACIENTE": dados.get("nome") or dados.get("NOME") or "Não informado",
                
                "DATA": agora.strftime("%d/%m/%Y"),
                "HORA": agora.strftime("%H:%M:%S"),
                "DATA_HORA": agora.strftime("%d/%m/%Y %H:%M:%S"),
                "DATA_PROCESSAMENTO": agora.strftime("%d/%m/%Y %H:%M:%S"),
                "TIMESTAMP": agora.isoformat(),
                "ARQUIVO_FONTE": "Webhook N8N Cloud"
            }
        
        return {
            "status": "success",
            "message": "Dados processados com sucesso",
            "dados": dados_extraidos,
            "timestamp": agora.isoformat(),
            "data_processamento": agora.strftime("%d/%m/%Y %H:%M:%S"),
            "environment": "production"
        }
        
    except Exception as e:
        logger.error(f"Erro no webhook: {e}")
        return {
            "status": "error",
            "message": f"Erro no processamento: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }

# Adicionar novo endpoint específico para Z-API
@app.post("/gerar-documento-zapi")
async def gerar_documento_zapi(request: MensagemRequest):
    """Endpoint específico para Z-API com formato exato que ela espera"""
    logger.info("=== GERAÇÃO DE DOCUMENTO PARA Z-API ===")
    logger.info(f"Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extrair dados
        dados_extraidos = extrair_dados_da_mensagem(request.mensagem)
        logger.info("Dados extraídos para Z-API")
        
        # Nome do arquivo simplificado
        nome_cliente = re.sub(r'[^\w]', '', dados_extraidos.get("NOME", "cliente"))[:10]
        timestamp = datetime.now().strftime('%d%m_%H%M')
        filename = f"{nome_cliente}_{timestamp}.docx"
        output_path = os.path.join(temp_dir, filename)
        
        # Criar documento
        try:
            possible_templates = ["template.docx", "modelo.docx", "templates/template.docx"]
            template_encontrado = None
            for template_path in possible_templates:
                if os.path.exists(template_path):
                    template_encontrado = template_path
                    break
            
            if template_encontrado:
                preencher_modelo(template_encontrado, output_path, dados_extraidos)
            else:
                criar_documento_fallback(dados_extraidos, output_path)
                
        except Exception as e:
            logger.error(f"Erro: {e}")
            criar_documento_fallback(dados_extraidos, output_path)
        
        # Verificar arquivo
        if not os.path.exists(output_path):
            raise Exception("Documento não foi gerado")
        
        file_size = os.path.getsize(output_path)
        if file_size < 1000:
            raise Exception("Arquivo muito pequeno - possível corrupção")
        
        # Ler e validar arquivo
        with open(output_path, "rb") as f:
            file_bytes = f.read()
            
        # Validar se é DOCX válido
        if not file_bytes.startswith(b'PK'):
            raise Exception("Arquivo não é um DOCX válido")
        
        # Gerar base64 sem quebras de linha
        base64_string = base64.b64encode(file_bytes).decode('ascii')
        
        # Validar base64
        if len(base64_string) < 1000:
            raise Exception("Base64 muito pequeno")
        
        # Testar se base64 pode ser decodificado
        try:
            base64.b64decode(base64_string)
        except Exception:
            raise Exception("Base64 inválido gerado")
        
        logger.info(f"✅ Arquivo: {filename} ({file_size} bytes)")
        logger.info(f"✅ Base64: {len(base64_string)} caracteres")
        
        return {
            "success": True,
            "filename": filename,
            "base64": base64_string,
            "size": file_size,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "dados": dados_extraidos,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"❌ ERRO: {e}")
        return {
            "success": False,
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }
    
    finally:
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        except:
            pass

@app.post("/test-docx")
async def test_docx():
    """Endpoint para testar geração de DOCX simples"""
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Criar documento de teste
        doc = Document()
        doc.add_heading('Teste de Documento', 0)
        doc.add_paragraph('Este é um teste de geração de DOCX.')
        doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        
        filename = f"teste_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        # Ler arquivo
        with open(filepath, "rb") as f:
            file_bytes = f.read()
        
        # Gerar base64
        base64_string = base64.b64encode(file_bytes).decode('ascii')
        
        return {
            "success": True,
            "message": "Teste OK",
            "filename": filename,
            "size": len(file_bytes),
            "base64_length": len(base64_string),
            "base64_preview": base64_string[:100] + "...",
            "is_valid_docx": file_bytes.startswith(b'PK'),
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }
    
    finally:
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        except:
            pass

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    logger.info("🚀 Iniciando servidor FastAPI para N8N + WhatsApp (Cloud Version)...")
    logger.info(f"🌐 Porta: {port}")
    logger.info(f"📅 Data atual: {datetime.now().strftime('%d/%m/%Y')}")
    logger.info(f"🕐 Hora atual: {datetime.now().strftime('%H:%M:%S')}")
    
    uvicorn.run(app, host="0.0.0.0", port=port) 