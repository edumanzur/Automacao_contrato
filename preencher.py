from docx import Document

def substituir_em_runs_preservando_tudo(paragrafos, dados):
    """
    Substitui placeholders nos runs preservando formatação
    """
    for paragrafo in paragrafos:
        # Primeiro, vamos concatenar todo o texto do parágrafo
        texto_completo = ""
        for run in paragrafo.runs:
            texto_completo += run.text
        
        # Verificar se há placeholders no texto completo
        texto_modificado = texto_completo
        substituicoes_feitas = False
        
        for chave, valor in dados.items():
            placeholder = f'{{{{{chave}}}}}'
            if placeholder in texto_modificado:
                # Garantir que valor é string
                valor_str = str(valor) if valor is not None else ""
                texto_modificado = texto_modificado.replace(placeholder, valor_str)
                substituicoes_feitas = True
                print(f"✅ Substituído: {placeholder} -> {valor_str}")
        
        # Se houve substituições, redistribuir o texto pelos runs
        if substituicoes_feitas:
            # Limpar todos os runs primeiro
            for run in paragrafo.runs:
                run.text = ""
            
            # Se não há runs, criar um
            if not paragrafo.runs:
                paragrafo.add_run()
            
            # Colocar todo o texto modificado no primeiro run
            # Isso preserva a formatação básica do primeiro run
            paragrafo.runs[0].text = texto_modificado

def preencher_modelo(caminho_modelo, caminho_saida, dados):
    """
    Preenche um modelo DOCX com os dados fornecidos
    
    Args:
        caminho_modelo (str): Caminho para o arquivo modelo .docx
        caminho_saida (str): Caminho onde salvar o arquivo preenchido
        dados (dict): Dicionário com os dados para substituição
    """
    try:
        print(f"📄 Abrindo modelo: {caminho_modelo}")
        doc = Document(caminho_modelo)
        
        # Converter todos os valores para string e tratar None
        dados_limpos = {}
        for chave, valor in dados.items():
            if valor is None or valor == "":
                dados_limpos[chave] = "Não informado"
            else:
                dados_limpos[chave] = str(valor)
        
        print(f"📝 Dados a serem substituídos:")
        for chave, valor in dados_limpos.items():
            valor_preview = valor[:50] + "..." if len(str(valor)) > 50 else valor
            print(f"   {{{{{chave}}}}} -> {valor_preview}")
        
        # Substituir nos parágrafos principais
        print("🔄 Processando parágrafos principais...")
        substituir_em_runs_preservando_tudo(doc.paragraphs, dados_limpos)
        
        # Substituir nas tabelas
        print("🔄 Processando tabelas...")
        for i, tabela in enumerate(doc.tables):
            print(f"   Tabela {i+1}...")
            for j, linha in enumerate(tabela.rows):
                for k, celula in enumerate(linha.cells):
                    substituir_em_runs_preservando_tudo(celula.paragraphs, dados_limpos)
        
        # Substituir nos cabeçalhos e rodapés
        print("🔄 Processando cabeçalhos e rodapés...")
        for i, section in enumerate(doc.sections):
            print(f"   Seção {i+1}...")
            
            # Cabeçalho
            if section.header:
                substituir_em_runs_preservando_tudo(section.header.paragraphs, dados_limpos)
            
            # Rodapé
            if section.footer:
                substituir_em_runs_preservando_tudo(section.footer.paragraphs, dados_limpos)
        
        # Salvar documento
        print(f"💾 Salvando documento em: {caminho_saida}")
        doc.save(caminho_saida)
        print(f"✅ Arquivo gerado com sucesso em: {caminho_saida}")
        
        return True
        
    except FileNotFoundError:
        error_msg = f"❌ Modelo não encontrado: {caminho_modelo}"
        print(error_msg)
        raise Exception(error_msg)
    
    except Exception as e:
        error_msg = f"❌ Erro ao preencher modelo: {str(e)}"
        print(error_msg)
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        raise Exception(error_msg)

# Função de teste para validar o preenchimento
def testar_preencher_modelo():
    """
    Função para testar o preenchimento com dados de exemplo
    """
    dados_teste = {
        "NOME": "Eduardo Silva",
        "EMAIL": "eduardo@gmail.com",
        "CPF": "123.456.789-00",
        "ENDERECO": "Rua das Flores, 123, Bloco A, Apt 45",
        "CEP": "12345-678",
        "TELEFONE": "(11) 99999-9999",
        "VALOR": "1.500,00",
        "PARCELAS": "12",
        "FORMA_PAGAMENTO": "Cartão de Crédito",
        "DATA_PROCESSAMENTO": "03/08/2025 14:30:00"
    }
    
    try:
        preencher_modelo("template.docx", "teste_output.docx", dados_teste)
        print("✅ Teste realizado com sucesso!")
        return True
    except Exception as e:
        print(f"❌ Erro no teste: {e}")
        return False

if __name__ == "__main__":
    # Executar teste se o arquivo for rodado diretamente
    testar_preencher_modelo()